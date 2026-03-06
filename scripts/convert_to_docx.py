import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(input_path, output_path):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)

    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found.")
        return

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Handle Headings
        if line.startswith('# '):
            if in_table:
                process_table(doc, table_data)
                in_table = False
                table_data = []
            heading = doc.add_heading(line[2:], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            if in_table:
                process_table(doc, table_data)
                in_table = False
                table_data = []
            doc.add_heading(line[3:], level=1)
            continue
        elif line.startswith('### '):
            if in_table:
                process_table(doc, table_data)
                in_table = False
                table_data = []
            doc.add_heading(line[4:], level=2)
            continue
        
        # Handle Tables (Simple Markdown Tables)
        if line.startswith('|') and '|' in line:
            if '---' in line:
                continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                in_table = True
                table_data.append(cells)
            continue
        else:
            if in_table:
                process_table(doc, table_data)
                in_table = False
                table_data = []
        
        # Handle normal text
        if line:
            paragraph = doc.add_paragraph(line)
        else:
            doc.add_paragraph()

    # Final table check
    if in_table:
        process_table(doc, table_data)

    doc.save(output_path)
    print(f"Successfully converted {input_path} to {output_path}")

def process_table(doc, table_data):
    if not table_data:
        return
    
    rows = len(table_data)
    cols = len(table_data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                table.cell(i, j).text = cell_text

if __name__ == "__main__":
    docs_to_convert = [
        "functional_test_design_and_implementation.md",
        "thesis_outline.md",
        "3.md"
    ]
    
    base_dir = r"d:\1下载\连锁零售门店库存管理系统 (1)\毕业论文文档"
    
    for filename in docs_to_convert:
        input_file = os.path.join(base_dir, filename)
        output_file = os.path.join(base_dir, filename.replace('.md', '.docx'))
        markdown_to_docx(input_file, output_file)
